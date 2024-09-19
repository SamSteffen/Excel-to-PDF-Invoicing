# import dependencies
import pandas as pd
import os
import openpyxl
import docx
import docx2pdf 

from os import remove, path
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx2pdf import convert
from datetime import datetime

# define the filepath to the desired xlsx workbook and docx document
cwd                = os.getcwd()
xlsx_filename      = 'Invoice_Data.xlsx'
xlsx_full_filepath = cwd + '\\' + xlsx_filename

# access the xlsx workbook
wb = openpyxl.load_workbook(filename = xlsx_full_filepath)

# create a dictionary of data from each worksheet
# 1. 'Business'
# create a dictionary of all the business owner data
business_header_list     = [row.value for row in wb['Business'][1]]
business_data_list       = [row.value for row in wb['Business'][2]]
business_data_dictionary = dict(zip(business_header_list, business_data_list))

# 2. 'Clients'
# retrieve the client numbers from the 'Clients' worksheet
client_count            = max([item.value for item in wb['Clients']['A'] if isinstance(item.value, int) == True])
extra                   = client_count + 1
client_numbers_list_int = list(range(0, extra))
client_numbers_list_str = [str(int) for int in client_numbers_list_int]

# create a dictionary of all the client data, using the CLIENT_NUMBERS as keys and the remainder of the client data as values
client_data_dictionary  = {}
# remove the first item in the list so the full list can be iterated (row '0' is a header column)
client_numbers_list_int.pop(0)
# add the extra item to the list so the full list can be iterated
client_numbers_list_int.append(extra)
for row_num in client_numbers_list_int:
    client_data_list = []
    for row in wb['Clients'][row_num]:
        client_data_list.append(row.value)
        client_data_dictionary[(row_num-1)] = client_data_list
        
# modify the dictionary to make the data readable/usable
for datarow in list(client_data_dictionary.values())[1:]:
    datarow[8] = datarow[8].strftime('%m-%d-%Y')

# 2. 'Timesheet'
timesheet_data_dictionary = {}
entry_count = len([item.value for item in wb['Timesheet']['A'] if isinstance(item.value, int) == True])
extra = entry_count + 1
timesheet_entry_rows_list = list(range(0, extra))

# create a dictionary of all the timesheet data, using ROW_NUMBERS as keys and data as values
timesheet_data_dictionary = {}
# remove the first item in the list so full list can be iterated (row '0' is a header column)
timesheet_entry_rows_list.pop(0)
# add the extra item to the list so the full list can be iterated
timesheet_entry_rows_list.append(extra)
for row_num in timesheet_entry_rows_list:
    timesheet_data_list = []
    for row in wb['Timesheet'][row_num]:
        timesheet_data_list.append(row.value)
        timesheet_data_dictionary[(row_num-1)] = timesheet_data_list

# modify the dictionary to make the data readable/usable
for data_row in list(timesheet_data_dictionary.values())[1:]:
    for client_data in list(client_data_dictionary.values()):
        if data_row[0] == client_data[0]:                           
            data_row[1] = client_data[1]   # client name
            data_row[2] = client_data[2]   # client address
            # Date of Service
            data_row[6] = datetime(data_row[5], data_row[4], data_row[3]).strftime('%m-%d-%Y')   # Date of Service
            data_row[7] = datetime(data_row[5], data_row[4], data_row[3]).strftime('%A')         # Weekday of Service
            data_row[8] = datetime(data_row[5], data_row[4], data_row[3]).strftime('%B')         # Month of Service
            data_row[9] = data_row[9].strftime('%H:%M')                                          # start time
            data_row[10] = data_row[10].strftime('%H:%M')                                        # end time
            data_row[11] = (datetime.strptime(data_row[10], '%H:%M')\
                            -datetime.strptime(data_row[9],'%H:%M')).seconds/3600                # hrs billed
            data_row[15] = ((datetime.strptime(data_row[10], '%H:%M')\
                            -datetime.strptime(data_row[9],'%H:%M')).seconds/3600)*data_row[14] # client per diem

# 3. 'Invoices'
# create a dictionary of all the invoice data, using the INVOICE_NUMBERS as keys and the remainder of the invoice data as values
# be sure to replace anything that appears as formulas with actual data from the client_data_list
invoice_count = len([item.value for item in wb['Invoices']['B']])-1 # subtract 1 because the first row is a header
extra = invoice_count + 1
invoice_list = [item.value for item in wb['Invoices']['B']]

invoice_data_dictionary = {}
invoice_list.pop(0)
invoice_list.insert(0, '00000')
invoice_data_rows_list = list(range(1, len(invoice_list)+1))

for i, invoice_number in enumerate(invoice_list):
    invoice_data_list = []
    for data in wb['Invoices'][i+1]:
        invoice_data_list.append(data.value)
        invoice_data_dictionary[invoice_number] = invoice_data_list
        
# create a unique list of the client_numbers that appear in the timesheet dataset
client_numbers_to_invoice_list = list(set([data_row[0] for data_row in list(timesheet_data_dictionary.values())[1:]]))

# modify the dictionary to make the data readable/usable
for data_row in list(invoice_data_dictionary.values())[1:]:
    for client_data in list(client_data_dictionary.values()):
        if data_row[5] == client_data[0]:
            data_row[2]  = data_row[2].strftime('%m-%d-%Y') # Invoice Date
            data_row[3]  = data_row[3].strftime('%m-%d-%Y') # Period Start Date
            data_row[4]  = data_row[4].strftime('%m-%d-%Y') # Period End Date
            data_row[6]  = client_data[1]                   # Client Name
            data_row[7]  = client_data[2]                   # Client Address
            data_row[8]  = client_data[3]                   # Client Phone (Primary)
            data_row[9]  = client_data[5]                   # Client Email (Primary)
            data_row[10] = client_data[7]                   # Preferred Payment Method
            data_row[11] = client_data[8]                   # Enrollment Date
            
    # create an empty list to hold valid values to sum
    hrs_to_invoice_values = []
    subtotal_values       = []
    
    # take each row of data from the timesheet wksheet...
    for timesheet_data in list(timesheet_data_dictionary.values()):
        # ensure that the data_row in the invoice_dictionary is meant for inclusion
        if data_row[5] in client_numbers_to_invoice_list:
            # sum all the values in the timesheet_data_dictionary
            # ensure that the data_row in the invoice_dictionary has the (1) same client_number as the timesheet item
            if data_row[5] == timesheet_data[0]:                
                # ensure that the data_row in the invoice_dictionary has 
                # (2) a date between the start and enddate on the timesheet data
                if data_row[3] <= timesheet_data[6] <= data_row[4]:
                    
                    # return the hours invoiced 
                    hrs_to_invoice_values.append(timesheet_data[11])
                    subtotal_values.append(timesheet_data[15])
        else:
            data_row[12] = 0
            data_row[13] = 0
            data_row[14] = 0

    # replace formulas with valid sums of data
    data_row[12] = sum(hrs_to_invoice_values)   # Hrs Invoiced
    data_row[13] = sum(subtotal_values)         # Subtotal
    data_row[14] = round(data_row[13]*0.05, 2)  # GST

# retrieve the headers from the data_dictionaries and store them as lists
client_data_header_list    = client_data_dictionary[0]
invoice_data_header_list   = invoice_data_dictionary['00000']
timesheet_data_header_list = timesheet_data_dictionary[0]

# create a list of dictionaries that associate the header with the data
individual_invoice_list = []
for data_row in list(invoice_data_dictionary.values())[1:]:    
    temp_dict = {}
    for j, value in enumerate(data_row):
        temp_dict[invoice_data_header_list[j]] = data_row[j]
    individual_invoice_list.append(temp_dict)

# 4. 'Descriptions'
# create a list to store distinct descriptions of service as keys in a forthcoming dictionary
distinct_service_descriptions_list = []

# loop through data and collect distinct instances of service descriptions
for i, data in enumerate(individual_invoice_list):        
    # iterate through each row of the timesheet_data_dictionary where the invoice number is the same as defined
    for row in timesheet_data_dictionary.values():
        # only target data for which invoice generation option is indicated
        if data['Generate Invoice?'] == 'Yes':        
            # match client number from individual_invoice_list to the client number in the timesheet_data_dictionary
            if data['Client Number'] == row[0]:
                # select only rows featuring work that occurred between desired period
                if data['Period Start Date'] <= row[6] <= data['Period End Date']:
                    if row[12] not in distinct_service_descriptions_list:
                        distinct_service_descriptions_list.append(row[12])

# create a new dictionary to store the data pertinent to service_descriptions
service_description_dictionary = {}
for service_description in distinct_service_descriptions_list:
    data_list = []
    for row in timesheet_data_dictionary.values():
        if row[12] == service_description:
            dictionary = {
                'Client Number'          : row[0]
                ,'Date of Service'       : row[6]
                ,'Hours'                 : row[11]
                ,'Description of Service': row[12]
                ,'Workers'               : row[13]
                ,'Rate/hr (CAD)'         : row[14]
                ,'Client Per Diem'       : row[15]
            }
            
            # add the extracted data to the data_list
            data_list.append(dictionary)
            
    service_description_dictionary[service_description] = data_list

# 5. Create a final_invoice_dict object that will be the final output printed on each invoice sheet
#create a final_invoice_dict to store outputs
final_invoice_dict = {}

# create a final_data_list to store outputs
final_data_list = []

# loop through the invoice_data and select only data rows for which invoices should be generated
for invoice_data_row in invoice_data_dictionary.values():
    if invoice_data_row[0] == "Yes":
        semifinal_data_list = []
        
        # loop through the service description dictionary and select data that matches to:
        # i   - invoice_data_client_number
        # ii  - invoice_data_start_date
        # iii - invoice_data_end_date
        for service_description, service_description_data in service_description_dictionary.items():
            for service_description_data_row in service_description_data:
                per_diem_data_list = []
                # if description dict client_number matches invoice_data_client_number
                if service_description_data_row['Client Number'] == invoice_data_row[5]:
                    # and description dict service_date is between invoice_data_start_date and invoice_data_end_date
                    if invoice_data_row[3] <= service_description_data_row['Date of Service'] <= invoice_data_row[4]:
                        # create a list to store the data needed for the final invoice
                        per_diem_data_list.append(service_description_data_row['Client Number'])
                        per_diem_data_list.append(service_description_data_row['Date of Service'])
                        per_diem_data_list.append(service_description)
                        per_diem_data_list.append(service_description_data_row['Workers'])
                        per_diem_data_list.append(service_description_data_row['Hours'])
                        per_diem_data_list.append(service_description_data_row['Rate/hr (CAD)'])
                        per_diem_data_list.append(service_description_data_row['Client Per Diem'])
                if per_diem_data_list != []:
                    semifinal_data_list.append(per_diem_data_list)
        
        if semifinal_data_list != []:
            # add the list of data to the final list
            final_data_list.append(semifinal_data_list)

for invoice_data in final_data_list:
    final_invoice_dict[invoice_data[0][0]] = invoice_data

# 6. Create a subtotal dict
# subtotal should add last item for each data row in final_invoice_dict
subtotal_dict = {}
for client_number, per_diem_data_list in final_invoice_dict.items():
    client_per_diem_amounts = []
    for data in per_diem_data_list:
        if data[0] == client_number:
            client_per_diem_amounts.append(data[6])
    subtotal = sum(client_per_diem_amounts)
    subtotal_dict[client_number] = subtotal

# 7. Create a GST dict
gst_dict = {}
for client_number, subtotal in subtotal_dict.items():
    gst_dict[client_number] = round(subtotal*0.05, 2)

# 8. Create a total dict
totals_dict = {}
for subtotal_client_number, subtotal in subtotal_dict.items():
    for gst_client_number, gst in gst_dict.items():
        if subtotal_client_number == gst_client_number:
            totals_dict[subtotal_client_number] = subtotal+gst

# create a function that will format the font, size and alignment of a paragraph object
def format_paragraph_obj(paragraph_obj, font, fontsize, alignment, space_before_pt=1, space_after_pt=1, style=None):
   # 0 = left, 1 = center, 2 = right, 3 = justify
    if alignment == 'right':
        paragraph_obj.alignment = 2
    elif alignment == 'centered':
        paragraph_obj.alignment = 1
    elif alignment == 'justified':
        paragraph_obj.alignment = 3
    else:
        paragraph_obj.alignment = 0
        
    obj_font = paragraph_obj.runs[0].font
    obj_font.name = font
    obj_font.size = Pt(fontsize)
    
    # space before text
    paragraph_obj.paragraph_format.space_before = Pt(space_before_pt)
    
    # space after text
    paragraph_obj.paragraph_format.space_after = Pt(space_after_pt)

    if style == 'Bold':
        paragraph_obj.runs[0].bold=True
    if style == 'Italics':
        paragraph_obj.runs[0].bold=True

# create a function that allows you to set the widths of columns (up to 4 columns)
# function assumes there will be at least 2 columns
def set_table_column_widths(table_obj, total_number_of_columns, width_1, width_2, width_3=0, width_4=0, width_5=0):
    width_list = [width_1, width_2, width_3, width_4, width_5]
    for i in range(0, total_number_of_columns):
        for cell in table_obj.columns[i].cells:
            cell.width = Inches(width_list[i])

# Create a function to set background shading for Header Rows
# https://stackoverflow.com/questions/26752856/python-docx-set-table-cell-background-and-text-color
def set_table_header_bg_color(table_obj, hex_color='ffffff'):
    tblCell = cell._tc
    tblCellProperties = tblCell.get_or_add_tcPr()
    clShading = OxmlElement('w:shd')
    # Hex of Dark Blue Shade {R:0x00, G:0x51, B:0x9E}
    # White = 'ffffff'
    # Black = '000000'
    # dark Blue = '351c75'
    clShading.set(qn('w:fill'), hex_color) 
    tblCellProperties.append(clShading)
    
    return cell

def format_column_cells(table_obj, column_number, alignment, space_before_pt, space_after_pt, font, fontsize, style='Normal'):
    for i, column in enumerate(table_obj.columns):
        if i == column_number:
            for cell in column.cells:
                
                for paragraph_obj in cell.paragraphs:
                    # set alignment
                    if alignment == 'right':
                        paragraph_obj.alignment = 2
                    elif alignment == 'centered':
                        paragraph_obj.alignment = 1
                    elif alignment == 'justified':
                        paragraph_obj.alignment = 3
                    else:
                        paragraph_obj.alignment = 0

                    # space before text
                    paragraph_obj.paragraph_format.space_before = Pt(space_before_pt)                        
                        
                    # space after text
                    paragraph_obj.paragraph_format.space_after = Pt(space_after_pt)

                    # set font
                    for run in paragraph_obj.runs:
                        obj_font = run.font
                        obj_font.name = font
                        obj_font.size = Pt(fontsize)
                        
                    # set style
                        if style == 'Bold':
                            run.bold=True
                        if style == 'Italics':
                            run.italics=True

def format_row_cells(table_obj, row_number, font, fontsize, alignment, space_before_pt=1, space_after_pt=1, style=None, cellcolor='ffffff'):
    for i, row in enumerate(table_obj.rows):
        if i == row_number:
            for cell in row.cells:
                
                # set cellcolor
                shade_obj = OxmlElement('w:shd')
                shade_obj.set(qn('w:fill'), cellcolor)
                cell._tc.get_or_add_tcPr().append(shade_obj)              
                
                for paragraph_obj in cell.paragraphs:
                    # set alignment
                    if alignment == 'right':
                        paragraph_obj.alignment = 2
                    elif alignment == 'centered':
                        paragraph_obj.alignment = 1
                    elif alignment == 'justified':
                        paragraph_obj.alignment = 3
                    else:
                        paragraph_obj.alignment = 0

                    # space before text
                    paragraph_obj.paragraph_format.space_before = Pt(space_before_pt)                        
                        
                    # space after text
                    paragraph_obj.paragraph_format.space_after = Pt(space_after_pt)
                        
                    # set font
                    obj_font = paragraph_obj.runs[0].font
                    obj_font.name  = font
                    obj_font.size  = Pt(fontsize)
#                     obj_font.color = textcolor

                    # set style
                    if style == 'Bold':
                        paragraph_obj.runs[0].bold=True
                    if style == 'Italics':
                        paragraph_obj.runs[0].italics=True

# loop through the invoice data to generate invoices for all clients                    
for i, data in enumerate(individual_invoice_list):
    # only target data for which invoice generation option is indicated
    if data['Generate Invoice?'] == 'Yes':
                
        #create a save string
        client_last_name = data['Client Name'].split(' ', -1)[-1]
        save_string = 'Invoice_'+data['Invoice Number']+'_'+client_last_name+'_'+data['Invoice Date']
        docx_save_string = save_string +'.docx'
        pdf_save_string  = 'PDFs/'+ save_string +'.pdf'

        if path.exists(docx_save_string):
            remove(docx_save_string)

        if path.exists(pdf_save_string):
            remove(pdf_save_string)

        # create a document object from the invoice template to store invoice information
        document = Document('Invoice_Template.docx')

        client_billing_info_list = list(individual_invoice_list[i].values())[6:10]

        # clear the document of unwanted lines
        # document._body.clear_content()

        ########################
        # CREATE THE HEADER TEXT
        ########################
        # Line 1 : 'Invoice'
        invoice_line = document.add_paragraph()
        invoice_line.add_run('Invoice')

        # FORMAT: Bahnschrift Light, 15 pt, right justified
        format_paragraph_obj(invoice_line, 'Bahnschrift Light', 15, 'right', 1, 1, 'Bold')            

        # Line 2 : 'Date'
        datestring = datetime.strptime(data['Invoice Date'], '%m-%d-%Y').strftime('%B %d, %Y')
        datestring_line = document.add_paragraph()
        datestring_line.add_run(datestring.upper())
        # FORMAT : Yu Gothic UI Semilight, 12 pt, right justified, ALL CAPS
        format_paragraph_obj(datestring_line, 'Yu Gothic UI Semilight', 12, 'right')

        # Line 3 : 'Invoice Number : '
        invoice_number = data['Invoice Number']
        invoice_number_line = document.add_paragraph()
        invoice_number_line.add_run(f'Invoice Number: {invoice_number}')
        # FORMAT : Yu Gothic UI Semilight, 12 pt, right justified
        format_paragraph_obj(invoice_number_line, 'Yu Gothic UI Semilight', 12, 'right')

        ##########################
        # CREATE THE TO/FROM TABLE
        ##########################
        # A 2-column table 

        # From:                          To:                 Bahnschrift Light, 13pt, left, Bold
        # Owner Name                     Client Name         Bahnschrift Light, 12pt, center
        # Owner Email                    Client Phone        Bahnschrift Light, 12pt, center
        # Owner Phone                                        Bahnschrift Light, 12pt, center

        #add a table of the description of services
        to_from_table = document.add_table(rows=4, cols=2)

        # create the table header rows by defining the cells
        header_cells              = to_from_table.rows[0].cells
        name_cells                = to_from_table.rows[1].cells
        contact_cells1            = to_from_table.rows[2].cells
        contact_cells2            = to_from_table.rows[3].cells

        # retrieve the data for the cells in the table
        to_from_table_header_list         = ['From: ','To: ']
        to_from_table_name_list           = [business_data_dictionary['Owner Name'], data['Client Name']]
        to_from_table_contact_cells1_list = [business_data_dictionary['Owner Email'], data['Client Phone (Primary)']]
        to_from_table_contact_cells2_list = [business_data_dictionary['Owner Phone'], '']

        # insert the data to their appropriate cells
        for i in list(range(0, len(to_from_table_header_list))):
            header_cells[i].text    = to_from_table_header_list[i]
            name_cells[i].text      = to_from_table_name_list[i]
            contact_cells1[i].text  = to_from_table_contact_cells1_list[i]
            contact_cells2[i].text  = to_from_table_contact_cells2_list[i]

        # format the table rows
        format_row_cells(to_from_table, 0, 'Bahnschrift Light', 13, 'left', 1, 1, 'Bold')            
        format_row_cells(to_from_table, 1, 'Bahnschrift Light', 12, 'left')            
        format_row_cells(to_from_table, 2, 'Bahnschrift Light', 12, 'left')            
        format_row_cells(to_from_table, 3, 'Bahnschrift Light', 12, 'left')            

        #################################### add a carriage return
        document.add_paragraph().add_run('\n')
        ####################################

        #######################################
        # CREATE THE PER DIEM HEADER TABLE ROW
        #######################################            
        # A 4-column table 
        #'GARDENING SERVICES FOR THE MONTH OF __ 2024', 'wORK HOURS', 'RATE', 'COST'
        # ALL CAPS, CENTERED, WHITE TEXT, CALIBRI, 14, DARK BLUE BACKGROUND

        header_table = document.add_table(rows=1, cols=4)

        # create row objects
        header_cells               = header_table.rows[0].cells
        startdate                  = data['Period Start Date']
        enddate                    = data['Period End Date']
        cell1_headerstring         = (f'Gardening services from \n{startdate} to {enddate}').upper()
        header_text_list           = [cell1_headerstring, 'WORK HOURS', 'RATE', 'COST']

        # insert the header names into the header cells
        for i in list(range(0, len(header_text_list))):
            header_cells[i].text = header_text_list[i]

        # format the header cell text
        format_row_cells(header_table, 0, 'Calibri', 14, 'centered', 1, 1, 'Bold')

        # set table header background color to dark blue hex# 351c75
        # center header text vertically within the cell
        for row in header_table.rows:
            for cell in row.cells:
                set_table_header_bg_color(cell, '351c75')
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # set table column widths
        set_table_column_widths(header_table, 4, 3.5, 2.0, 1.0, 1.0)

        # set table row height
        for row in header_table.rows:
            row.height = Cm(0.7)

        ###########################
        # CREATE THE PER DIEM TABLE ROWS
        ###########################
        per_diem_table = document.add_table(rows=1, cols=4)

        # create empty lists for desired data
        description_data = []
        workers_qty_data = []
        work_hrs_data    = [] 
        rate_data        = []
        cost_data        = [] 

        for invoice_data_list in final_invoice_dict.values():
            for invoice_data in invoice_data_list: 
                # match the invoice data on the client number
                if invoice_data[0] == data['Client Number']:
        #           description_data.append(invoice_data[2])     # description

                    # only add the description once; if it's already there, add empty string ('')
                    if invoice_data[2] not in description_data:
                        description_data.append(invoice_data[2])     # description
                    else:
                        description_data.append('')

                    workers_qty_data.append(invoice_data[3])                   # workers_qty
                    work_hrs_data.append(invoice_data[4])                      # work_hrs
                    rate_data.append(format(invoice_data[5],'.2f'))                          # rate
                    cost_data.append(format(invoice_data[6],'.2f'))            # cost

        # convert ints and floats into strings
        # combine the workers and hours data into one cell
        workers_qty_and_hrs_data = [
                                    '('+str(workers_qty)+' ppl.) ' + str(work_hrs_data[i]) + 'h' \
                                    if workers_qty > 1 \
                                    else '('+str(workers_qty)+' pers.) ' + str(work_hrs_data[i]) + 'h' \
                                    for i, workers_qty in enumerate(workers_qty_data)\
                                   ]
        # workers_qty_data         = ['(' + str(data) + ' ppl.)' for data in workers_qty_data]
        # work_hrs_data            = [str(data) + ' h' for data in work_hrs_data]
        rate_data                = ['$ '+ str(data) + '/hr' for data in rate_data]
        cost_data                = [data +' $' for data in cost_data]

        for i in list(range(0, len(description_data))):
            row_cells = per_diem_table.add_row().cells
            row_cells[0].text = description_data[i]
            row_cells[1].text = workers_qty_and_hrs_data[i]
            row_cells[2].text = rate_data[i]
            row_cells[3].text = cost_data[i]

            # set table column widths
            set_table_column_widths(per_diem_table, 4, 3.5, 2.0, 1.0, 1.0)

        # format table columns
        format_column_cells(per_diem_table, 0, 'centered', 1, 1, 'Yu Gothic UI Semilight', 11)
        format_column_cells(per_diem_table, 1, 'justified', 1, 1, 'Yu Gothic UI Semilight', 13)
        format_column_cells(per_diem_table, 2, 'centered', 1, 1, 'Yu Gothic UI Semilight', 13, 'Italics')
        format_column_cells(per_diem_table, 3, 'right', 1, 1, 'Yu Gothic UI Semilight', 13)

        # set table row height
        for row in per_diem_table.rows:
            row.height = Cm(0.25)

        # add a solid line
        document.add_paragraph('------------------------------------------------------------------------------------------------------------------------------------------')

        ###########################
        # CREATE THE SUBTOTAL ROW
        ###########################
        # create new table for subtotal data
        subtotal_table_row = document.add_table(rows=1, cols=4)

        # create text for subtotal_table_row
        for subtotal_client_number, subtotal in subtotal_dict.items():
            if subtotal_client_number == data['Client Number']:
                subtotal_table_row_text_list = ['SUBTOTAL','','', format(subtotal, '.2f') + ' $']
            if data['Client Number'] not in list(subtotal_dict.keys()):
                subtotal_table_row_text_list = ['SUBTOTAL','','', '0.00 $']
                        
        subtotal_row_cells               = subtotal_table_row.rows[0].cells

        # insert the header names into the header cells
        for i in list(range(0, len(subtotal_table_row_text_list))):
            subtotal_row_cells[i].text = subtotal_table_row_text_list[i]

        # format subtotal row
        format_row_cells(subtotal_table_row, 0, 'Yu Gothic UI Semilight', 13, 'centered')

        # set table column widths
        set_table_column_widths(subtotal_table_row, 4, 3.5, 2.0, 1.0, 1.0)

        # set table row height
        for row in subtotal_table_row.rows:
            row.height = Cm(0.25)        

        ###########################
        # CREATE THE GST ROW
        ###########################
        # create new table for gst data
        gst_table_row = document.add_table(rows=1, cols=4)

        # create text for subtotal_table_row
        for gst_client_number, gst in gst_dict.items():
            if gst_client_number == data['Client Number']:
                gst_table_row_text_list = ['GST','5 %','', format(gst,'.2f') + ' $']
            if data['Client Number'] not in list(gst_dict.keys()):
                gst_table_row_text_list = ['GST','5 %','', '0.00 $']

        # create cell objects to hold text
        gst_row_cells               = gst_table_row.rows[0].cells

        # insert the header names into the header cells
        for i in list(range(0, len(gst_table_row_text_list))):
            gst_row_cells[i].text = gst_table_row_text_list[i]

        # format subtotal row
        format_row_cells(gst_table_row, 0, 'Yu Gothic UI Semilight', 11, 'right')

        # set table column widths
        set_table_column_widths(gst_table_row, 4, 3.5, 2.0, 1.0, 1.0)

        # set table row height
        for row in gst_table_row.rows:
            row.height = Cm(0.25)        

        # add a solid line
        document.add_paragraph('------------------------------------------------------------------------------------------------------------------------------------------')

        ###########################
        # CREATE THE TOTALS ROW
        ###########################
        # create new table for totals data
        totals_table_row = document.add_table(rows=1, cols=4)

        # create text for subtotal_table_row
        for totals_client_number, totals in totals_dict.items():
            if totals_client_number == data['Client Number']:
                totals_table_row_text_list = ['TOTAL','','', '$ ' + format(totals,'.2f')]
            if data['Client Number'] not in list(totals_dict.keys()):
                totals_table_row_text_list = ['TOTAL','','', '0.00 $']

        # create cell objects to hold text
        totals_row_cells               = totals_table_row.rows[0].cells

        # insert the header names into the header cells
        for i in list(range(0, len(totals_table_row_text_list))):
            totals_row_cells[i].text = totals_table_row_text_list[i]

        # format subtotal row
        format_row_cells(totals_table_row, 0, 'Yu Gothic UI Semilight', 13, 'left', 1, 1, 'Bold')

        # set table column widths
        set_table_column_widths(totals_table_row, 4, 3.5, 2.0, 1.0, 1.0)

        # set table row height
        for row in totals_table_row.rows:
            row.height = Cm(0.25)        

        # add a solid line
        document.add_paragraph('------------------------------------------------------------------------------------------------------------------------------------------')

        # Add 'To be paid upon receipt line'
        footer_line1 = document.add_paragraph()
        footer_line1.add_run('To be paid upon receipt.')
        # FORMAT: Yu Gothic UI Semilight, 12 pt, centered
        format_paragraph_obj(footer_line1, 'Yu Gothic UI Semilight', 12, 'centered')

        ####################################
        document.add_paragraph().add_run('\n')
        ####################################

        # Add 'Please do not hesitate to let me know if you have questions or concerns', centered
        footer_line2 = document.add_paragraph()
        footer_line2.add_run('Please do not hesitate to let me know if you have questions or concerns.')
        # FORMAT : Yu Gothic UI Semilight, 11 pt, centered
        format_paragraph_obj(footer_line2, 'Yu Gothic UI Semilight', 11, 'centered')            

        ####################################
        document.add_paragraph().add_run('\n')
        ####################################

        # Add 'Thank you!' line, centered
        footer_line3 = document.add_paragraph()
        footer_line3.add_run('Thank you!')
        # FORMAT : Yu Gothic UI Semilight, 14 pt, centered
        format_paragraph_obj(footer_line3, 'Yu Gothic UI Semilight', 14, 'centered')            

        # save the document as a word doc
        document.save(docx_save_string)

        # convert the word doc to a pdf, store it in a separate folder
        convert(docx_save_string, pdf_save_string)

        # delete the word doc version
        remove(docx_save_string)



